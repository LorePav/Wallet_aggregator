# ============================================================
# TOKENOMICS_EXPORTER.PY
# Genera report Word (.docx) e PDF (.pdf) per l'analisi
# tokenomics di asset crypto.
# Speculare a fundamental_exporter.py ma adattato a crypto.
# ============================================================

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

from crypto_fetcher import safe_nested, format_large_number
from crypto_analyzer import (
    get_valuation_metrics, get_tokenomics_metrics,
    get_onchain_metrics, get_developer_metrics,
    get_community_metrics,
)

# Colore accent crypto (verde/teal invece del blu delle azioni)
_ACCENT = (0x00, 0xA6, 0xB2)


# ============================================================
# WORD (.docx)
# ============================================================

def export_to_word(data: dict) -> bytes:
    name   = data.get("name", data.get("id", "N/D"))
    symbol = data.get("symbol", "").upper()
    oggi   = datetime.now().strftime("%d/%m/%Y %H:%M")
    desc   = data.get("description", {}).get("en", "Descrizione non disponibile.")
    rank   = data.get("market_cap_rank")
    mc     = safe_nested(data.get("market_data", {}), "market_cap", "usd")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Titolo
    title = doc.add_heading(f"Analisi Tokenomics: {name} ({symbol})", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(*_ACCENT)

    sub = doc.add_paragraph(f"Rank #{rank} | Market Cap: {format_large_number(mc)} | Generato il: {oggi}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # Descrizione
    _word_heading(doc, "1. Overview del Progetto")
    cats = ", ".join(data.get("categories", [])[:4]) or "N/D"
    _word_info_row(doc, "Categorie", cats)
    _word_info_row(doc, "Data Genesi", data.get("genesis_date", "N/D"))
    _word_info_row(doc, "Sito Web", (data.get("links", {}).get("homepage") or ["N/D"])[0])
    doc.add_paragraph()
    _word_heading2(doc, "Descrizione")
    clean_desc = str(desc).replace("<p>", "").replace("</p>", " ").replace("<br>", " ")[:1500]
    doc.add_paragraph(clean_desc)
    doc.add_paragraph()

    # Sezioni metriche
    _word_heading(doc, "2. Valutazione")
    _word_metrics_table(doc, get_valuation_metrics(data))

    _word_heading(doc, "3. Tokenomics")
    _word_metrics_table(doc, get_tokenomics_metrics(data))

    _word_heading(doc, "4. On-Chain & Mercato")
    _word_metrics_table(doc, get_onchain_metrics(data))

    _word_heading(doc, "5. Developer Activity")
    _word_metrics_table(doc, get_developer_metrics(data))

    _word_heading(doc, "6. Community")
    _word_metrics_table(doc, get_community_metrics(data))

    # Disclaimer
    doc.add_paragraph()
    disc = doc.add_paragraph(
        "⚠️ DISCLAIMER: Analisi generata automaticamente a scopo informativo. "
        "Non costituisce consulenza finanziaria. Il mercato crypto è altamente volatile. "
        "Dati da CoinGecko API."
    )
    disc.runs[0].font.size = Pt(8)
    disc.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _word_heading(doc, text: str):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(*_ACCENT)


def _word_heading2(doc, text: str):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.size = Pt(12)


def _word_info_row(doc, label: str, value: str):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    p.add_run(str(value)).font.size = Pt(10)


def _word_metrics_table(doc, metrics: dict):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metrica"
    hdr[1].text = "Valore"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for label, value in metrics.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _fmt(label, value)
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def _fmt(label: str, value) -> str:
    """Formatta un valore per l'export."""
    if value is None:
        return "N/D"
    if isinstance(value, str):
        return value
    if "%" in label and isinstance(value, float):
        return f"{value:.2f}%"
    if isinstance(value, (int, float)) and abs(value) > 1_000_000:
        return format_large_number(value)
    if isinstance(value, float):
        return f"{value:.4f}" if value < 0.01 else f"{value:.2f}"
    return str(value)


# ============================================================
# PDF (.pdf)
# ============================================================

class TokenomicsPDF(FPDF):
    def __init__(self, name: str, symbol: str):
        super().__init__()
        self.coin_name   = name
        self.coin_symbol = symbol

    def header(self):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*_ACCENT)
        self.cell(0, 8, f"Analisi Tokenomics: {self.coin_name} ({self.coin_symbol})", align="L")
        self.set_text_color(150, 150, 150)
        self.set_font("Helvetica", "", 8)
        self.cell(0, 8, datetime.now().strftime("%d/%m/%Y"), align="R",
                  new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-20)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(2)
        self.set_font("Helvetica", "I", 7)
        self.set_text_color(150, 150, 150)
        self.cell(0, 5,
                  f"Pagina {self.page_no()} | Dati: CoinGecko API | Solo scopo informativo, non è consulenza finanziaria.",
                  align="C")


def export_to_pdf(data: dict) -> bytes:
    name   = data.get("name", data.get("id", "N/D"))
    symbol = data.get("symbol", "").upper()
    rank   = data.get("market_cap_rank")
    mc     = safe_nested(data.get("market_data", {}), "market_cap", "usd")
    cats   = ", ".join(data.get("categories", [])[:3]) or "N/D"
    desc   = data.get("description", {}).get("en", "")
    desc_clean = str(desc).replace("<p>", "").replace("</p>", " ").replace("<br>", " ")[:500]

    pdf = TokenomicsPDF(name=name, symbol=symbol)
    pdf.set_margins(15, 15, 15)
    pdf.add_page()

    # Header
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(*_ACCENT)
    pdf.cell(0, 12, "Analisi Tokenomics", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, f"{name} ({symbol})", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Rank #{rank} | Market Cap: {format_large_number(mc)} | Categorie: {cats}",
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    # Descrizione
    _pdf_section(pdf, "1. Overview del Progetto")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(50, 50, 50)
    if desc_clean:
        pdf.multi_cell(0, 5, desc_clean + ("..." if len(str(desc)) > 500 else ""))
    pdf.ln(4)

    _pdf_section(pdf, "2. Valutazione")
    _pdf_table(pdf, get_valuation_metrics(data))

    _pdf_section(pdf, "3. Tokenomics")
    _pdf_table(pdf, get_tokenomics_metrics(data))

    _pdf_section(pdf, "4. On-Chain & Mercato")
    _pdf_table(pdf, get_onchain_metrics(data))

    _pdf_section(pdf, "5. Developer Activity")
    _pdf_table(pdf, get_developer_metrics(data))

    _pdf_section(pdf, "6. Community")
    _pdf_table(pdf, get_community_metrics(data))

    return bytes(pdf.output())


def _pdf_section(pdf: FPDF, title: str):
    pdf.set_fill_color(*_ACCENT)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, title, fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_text_color(50, 50, 50)


def _pdf_table(pdf: FPDF, metrics: dict):
    col_w_label, col_w_value, row_h = 100, 75, 6
    pdf.set_fill_color(220, 245, 245)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(col_w_label, row_h, "Metrica", border=1, fill=True)
    pdf.cell(col_w_value, row_h, "Valore",  border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    fills = [(255, 255, 255), (240, 253, 253)]
    for i, (label, value) in enumerate(metrics.items()):
        r, g, b = fills[i % 2]
        pdf.set_fill_color(r, g, b)
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(col_w_label, row_h, label[:45],         border=1, fill=True)
        pdf.cell(col_w_value, row_h, _fmt(label, value), border=1, fill=True,
                 new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

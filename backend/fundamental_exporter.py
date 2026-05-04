# ============================================================
# FUNDAMENTAL_EXPORTER.PY
# Genera report Word (.docx) e PDF (.pdf) con l'analisi
# fondamentale completa di un'azienda.
# Usato dagli endpoint /api/fundamental/{ticker}/export/
# ============================================================

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

from data_fetcher import safe_get, format_large_number
from analyzer import (
    get_valuation_metrics, get_profitability_metrics,
    get_growth_metrics, get_financial_health_metrics,
    get_dividend_metrics, get_market_metrics
)


# ============================================================
# ESPORTAZIONE IN WORD (.docx)
# ============================================================

def export_to_word(data: dict) -> bytes:
    info = data["info"]
    nome_azienda = safe_get(info, "shortName", data["ticker"])
    settore = safe_get(info, "sector", "N/D")
    industria = safe_get(info, "industry", "N/D")
    descrizione = safe_get(info, "longBusinessSummary", "Descrizione non disponibile.")
    oggi = datetime.now().strftime("%d/%m/%Y %H:%M")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading(f"Analisi Fondamentale: {nome_azienda}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0x1A, 0x53, 0xFF)

    sub = doc.add_paragraph(f"Ticker: {data['ticker']} | Generato il: {oggi}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    _word_heading(doc, "1. Panoramica Aziendale")
    _word_info_row(doc, "Settore", settore)
    _word_info_row(doc, "Industria", industria)
    _word_info_row(doc, "Market Cap", format_large_number(safe_get(info, "marketCap")))
    emp = safe_get(info, "fullTimeEmployees")
    _word_info_row(doc, "Dipendenti", f"{emp:,}" if emp else "N/D")
    _word_info_row(doc, "Paese", safe_get(info, "country", "N/D"))
    _word_info_row(doc, "Sito web", safe_get(info, "website", "N/D"))
    doc.add_paragraph()
    _word_heading2(doc, "Descrizione Aziendale")
    doc.add_paragraph(str(descrizione)[:1500])
    doc.add_paragraph()

    _word_heading(doc, "2. Dati di Mercato")
    _word_metrics_table(doc, get_market_metrics(info))

    _word_heading(doc, "3. Valutazione (Multipli di Mercato)")
    _word_metrics_table(doc, get_valuation_metrics(info))

    _word_heading(doc, "4. Redditività")
    _word_metrics_table(doc, get_profitability_metrics(info))

    _word_heading(doc, "5. Crescita")
    _word_metrics_table(doc, get_growth_metrics(info))

    _word_heading(doc, "6. Salute Finanziaria")
    _word_metrics_table(doc, get_financial_health_metrics(info))

    _word_heading(doc, "7. Dividendi")
    _word_metrics_table(doc, get_dividend_metrics(info))

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "⚠️ DISCLAIMER: Questa analisi è generata automaticamente a scopo informativo. "
        "Non costituisce consulenza finanziaria o di investimento. "
        "I dati provengono da Yahoo Finance. "
        "Prima di investire, consulta un professionista abilitato."
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _word_heading(doc, text: str):
    h = doc.add_heading(text, level=1)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x53, 0xFF)


def _word_heading2(doc, text: str):
    h = doc.add_heading(text, level=2)
    h.runs[0].font.size = Pt(12)


def _word_info_row(doc, label: str, value: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = p.add_run(str(value))
    run_value.font.size = Pt(10)


def _word_metrics_table(doc, metrics: dict):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metrica"
    hdr_cells[1].text = "Valore"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for label, value in metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = label
        if value is None:
            row_cells[1].text = "N/D"
        elif isinstance(value, float):
            if -1 <= value <= 1 and any(kw in label.lower() for kw in
                    ["margin", "yield", "roe", "roa", "payout", "crescita"]):
                row_cells[1].text = f"{value * 100:.2f}%"
            elif abs(value) > 1_000_000:
                row_cells[1].text = format_large_number(value)
            else:
                row_cells[1].text = f"{value:.2f}"
        elif isinstance(value, int):
            if abs(value) > 1_000_000:
                row_cells[1].text = format_large_number(value)
            else:
                row_cells[1].text = str(value)
        else:
            row_cells[1].text = str(value)

        for cell in row_cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()


# ============================================================
# ESPORTAZIONE IN PDF (.pdf)
# ============================================================

class StockPDF(FPDF):
    def __init__(self, company_name: str, ticker: str):
        super().__init__()
        self.company_name = company_name
        self.ticker = ticker

    def header(self):
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(26, 83, 255)
        self.cell(0, 8, f"Analisi Fondamentale: {self.company_name} ({self.ticker})", align="L")
        self.set_text_color(150, 150, 150)
        self.set_font("Helvetica", "", 8)
        oggi = datetime.now().strftime("%d/%m/%Y")
        self.cell(0, 8, f"Generato il {oggi}", align="R", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-20)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(2)
        self.set_font("Helvetica", "I", 7)
        self.set_text_color(150, 150, 150)
        self.cell(
            0, 5,
            f"Pagina {self.page_no()} | Dati: Yahoo Finance | Solo scopo informativo, non è consulenza finanziaria.",
            align="C"
        )


def export_to_pdf(data: dict) -> bytes:
    info = data["info"]
    nome_azienda = safe_get(info, "shortName", data["ticker"])
    settore = safe_get(info, "sector", "N/D")
    industria = safe_get(info, "industry", "N/D")
    descrizione = safe_get(info, "longBusinessSummary", "Descrizione non disponibile.")

    pdf = StockPDF(company_name=nome_azienda, ticker=data["ticker"])
    pdf.set_margins(15, 15, 15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(26, 83, 255)
    pdf.cell(0, 12, "Analisi Fondamentale", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, nome_azienda, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Ticker: {data['ticker']} | Settore: {settore} | Industria: {industria}",
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    _pdf_section_title(pdf, "1. Panoramica Aziendale")
    desc_trunc = str(descrizione)[:600] + ("..." if len(str(descrizione)) > 600 else "")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(50, 50, 50)
    pdf.multi_cell(0, 5, desc_trunc)
    pdf.ln(4)

    _pdf_section_title(pdf, "2. Dati di Mercato")
    _pdf_metrics_table(pdf, get_market_metrics(info))

    _pdf_section_title(pdf, "3. Valutazione (Multipli di Mercato)")
    _pdf_metrics_table(pdf, get_valuation_metrics(info))

    _pdf_section_title(pdf, "4. Redditività")
    _pdf_metrics_table(pdf, get_profitability_metrics(info))

    _pdf_section_title(pdf, "5. Crescita")
    _pdf_metrics_table(pdf, get_growth_metrics(info))

    _pdf_section_title(pdf, "6. Salute Finanziaria")
    _pdf_metrics_table(pdf, get_financial_health_metrics(info))

    _pdf_section_title(pdf, "7. Dividendi")
    _pdf_metrics_table(pdf, get_dividend_metrics(info))

    return bytes(pdf.output())


def _pdf_section_title(pdf: FPDF, title: str):
    pdf.set_fill_color(26, 83, 255)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, title, fill=True, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_text_color(50, 50, 50)


def _pdf_metrics_table(pdf: FPDF, metrics: dict):
    col_w_label = 100
    col_w_value = 75
    row_h = 6

    pdf.set_fill_color(230, 235, 255)
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(col_w_label, row_h, "Metrica", border=1, fill=True)
    pdf.cell(col_w_value, row_h, "Valore", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    fill_colors = [(255, 255, 255), (245, 247, 255)]
    for i, (label, value) in enumerate(metrics.items()):
        r, g, b = fill_colors[i % 2]
        pdf.set_fill_color(r, g, b)
        pdf.set_font("Helvetica", "", 9)

        if value is None:
            formatted = "N/D"
        elif isinstance(value, float):
            if -1 <= value <= 1 and any(kw in label.lower() for kw in
                    ["margin", "yield", "roe", "roa", "payout", "crescita"]):
                formatted = f"{value * 100:.2f}%"
            elif abs(value) > 1_000_000:
                formatted = format_large_number(value)
            else:
                formatted = f"{value:.2f}"
        elif isinstance(value, int) and abs(value) > 1_000_000:
            formatted = format_large_number(value)
        else:
            formatted = str(value)

        pdf.cell(col_w_label, row_h, label[:45], border=1, fill=True)
        pdf.cell(col_w_value, row_h, formatted, border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    pdf.ln(5)
